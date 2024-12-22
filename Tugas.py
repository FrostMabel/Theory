import os
from docx import Document
import docx
from docx.shared import Cm
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn


FILE_NAME = os.path.join(os.getcwd(), "pesanan.docx")

# Inisialisasi file jika belum ada
def init_file():
    if not os.path.exists(FILE_NAME):
        doc = Document()
        doc.add_heading("Data Pesanan Restoran", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = "Nama Pelanggan"
        header_cells[1].text = "Pesanan"
        header_cells[2].text = "Status"
        header_cells[3].text = "Gambar"
        doc.save(FILE_NAME)

# Fungsi untuk membaca semua pesanan
def read_pesanan():
    doc = Document(FILE_NAME)
    table = doc.tables[0]
    print("\nData Pesanan:")
    for row in table.rows[1:]:
        print(f"Nama Pelanggan: {row.cells[0].text}, Pesanan: {row.cells[1].text}, "
              f"Status: {row.cells[2].text}, Path Gambar: {row.cells[3].text}")

# Fungsi untuk menambahkan pesanan baru
def create_pesanan(nama, pesanan, path_gambar):
    doc = Document(FILE_NAME)
    table = doc.tables[0]
    row = table.add_row().cells
    row[0].text = nama
    row[1].text = pesanan
    row[2].text = "Diproses"

    # Add image to the cell if the path is valid
    if os.path.exists(path_gambar):
        paragraph = row[3].add_paragraph()
        run = paragraph.add_run()
        try:
            run.add_picture(path_gambar, width=docx.shared.Cm(3), height=docx.shared.Cm(3))  # Adjust image size
        except Exception as e:
            print(f"Gagal menambahkan gambar: {e}")
    else:
        row[3].text = "Path gambar tidak valid."
    doc.save(FILE_NAME)
    print("Pesanan berhasil ditambahkan.")


# Fungsi untuk memperbarui status pesanan
def update_pesanan(nama, status_baru):
    doc = Document(FILE_NAME)
    table = doc.tables[0]
    for row in table.rows[1:]:
        if row.cells[0].text == nama:
            row.cells[2].text = status_baru
            doc.save(FILE_NAME)
            print("Status pesanan berhasil diperbarui.")
            return
    print("Pesanan tidak ditemukan.")

# Fungsi untuk menghapus pesanan yang dibatalkan
def delete_pesanan(nama):
    doc = Document(FILE_NAME)
    old_table = doc.tables[0]

    # Collect header and rows to keep
    rows_to_keep = []
    for i, row in enumerate(old_table.rows[1:], start=1):  # Skip the header row
        if not (row.cells[0].text == nama and row.cells[2].text.lower() == "batal"):
            rows_to_keep.append(row)

    # Remove old table
    for _ in range(len(old_table.rows)):
        old_table._rows[0]._element.getparent().remove(old_table._rows[0]._element)

    # Add a new table
    new_table = doc.add_table(rows=1, cols=4)
    new_table.style = 'Table Grid'
    new_table.rows[0].cells[0].text = "Nama Pelanggan"
    new_table.rows[0].cells[1].text = "Pesanan"
    new_table.rows[0].cells[2].text = "Status"
    new_table.rows[0].cells[3].text = "Path Gambar"

    # Copy rows and images
    for row in rows_to_keep:
        new_row = new_table.add_row().cells
        for col_idx in range(3):  # Copy text for first three columns
            new_row[col_idx].text = row.cells[col_idx].text

        # Copy image in the fourth column if it exists
        image_cell = row.cells[3]
        if image_cell._element.xpath(".//w:drawing"):
            # Extract the image XML element
            drawing = image_cell._element.xpath(".//w:drawing")[0]
            new_row[3]._element.append(parse_xml(drawing.xml))

    doc.save(FILE_NAME)
    print("Pesanan berhasil dihapus (jika ditemukan dan statusnya 'Batal').")


# Fungsi untuk mencari pesanan berdasarkan nama pelanggan
def search_pesanan(nama):
    doc = Document(FILE_NAME)
    table = doc.tables[0]
    print("\nHasil Pencarian:")
    for row in table.rows[1:]:
        if nama.lower() in row.cells[0].text.lower():
            print(f"Nama Pelanggan: {row.cells[0].text}, Pesanan: {row.cells[1].text}, "
                  f"Status: {row.cells[2].text}, Path Gambar: {row.cells[3].text}")

# Menu utama
def main():
    init_file()
    while True:
        print("\nMenu:")
        print("1. Tambah Pesanan")
        print("2. Lihat Semua Pesanan")
        print("3. Perbarui Status Pesanan")
        print("4. Hapus Pesanan (Status Batal)")
        print("5. Cari Pesanan")
        print("6. Keluar")
        choice = input("Pilih menu: ")

        if choice == "1":
            nama = input("Masukkan nama pelanggan: ")
            pesanan = input("Masukkan pesanan: ")
            path_gambar = input("Masukkan path gambar: ")
            create_pesanan(nama, pesanan, path_gambar)
        elif choice == "2":
            read_pesanan()
        elif choice == "3":
            nama = input("Masukkan nama pelanggan yang akan diperbarui: ")
            status_baru = input("Masukkan status baru (Diproses/Selesai/Batal): ")
            update_pesanan(nama, status_baru)
        elif choice == "4":
            nama = input("Masukkan nama pelanggan yang akan dihapus: ")
            delete_pesanan(nama)
        elif choice == "5":
            nama = input("Masukkan nama pelanggan yang dicari: ")
            search_pesanan(nama)
        elif choice == "6":
            print("Keluar dari program.")
            break
        else:
            print("Pilihan tidak valid. Coba lagi.")

if __name__ == "__main__":
    main()
