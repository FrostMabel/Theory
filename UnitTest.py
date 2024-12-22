import os
import unittest
from docx import Document
from Tugas import init_file, read_pesanan, create_pesanan, update_pesanan, delete_pesanan, search_pesanan, FILE_NAME

class TestPesanan(unittest.TestCase):
    
    def setUp(self):
        # Create the initial file setup before each test
        init_file()
    
    def tearDown(self):
        # Clean up by removing the file after each test
        if os.path.exists(FILE_NAME):
            os.remove(FILE_NAME)

    def test_init_file(self):
        self.assertTrue(os.path.exists(FILE_NAME))
        doc = Document(FILE_NAME)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)  # Only header row should exist
        self.assertEqual(len(table.columns), 4)  # Four columns should exist
    
    def test_create_pesanan(self):
        create_pesanan("John Doe", "Pizza", "invalid_path.jpg")
        doc = Document(FILE_NAME)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 2)  # One header row and one data row
        self.assertEqual(table.cell(1, 0).text, "John Doe")
        self.assertEqual(table.cell(1, 1).text, "Pizza")
        self.assertEqual(table.cell(1, 2).text, "Diproses")
        self.assertEqual(table.cell(1, 3).text, "Path gambar tidak valid.")
    
    def test_update_pesanan(self):
        create_pesanan("John Doe", "Pizza", "invalid_path.jpg")
        update_pesanan("John Doe", "Selesai")
        doc = Document(FILE_NAME)
        table = doc.tables[0]
        self.assertEqual(table.cell(1, 2).text, "Selesai")
    
    def test_delete_pesanan(self):
        create_pesanan("John Doe", "Pizza", "invalid_path.jpg")
        update_pesanan("John Doe", "Batal")
        delete_pesanan("John Doe")
        doc = Document(FILE_NAME)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)  # Only header row should remain
    
    def test_search_pesanan(self):
        create_pesanan("John Doe", "Pizza", "invalid_path.jpg")
        create_pesanan("Jane Doe", "Burger", "invalid_path.jpg")
        results = search_pesanan("Doe")
        self.assertIn("John Doe", results)
        self.assertIn("Jane Doe", results)

if __name__ == '__main__':
    unittest.main()
